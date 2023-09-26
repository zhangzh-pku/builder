from typing import List
from models.base.dataset import Dataset, Document
from langchain.text_splitter import CharacterTextSplitter
from abc import ABC, abstractmethod
import io
import sys
from docx import Document as WordDocument
from models.data_loader.options import PDFSplitterOption, PDFEmbeddingOption, PDFRetrivalOption
from langchain.schema import Document
from loguru import logger
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfinterp import PDFPageInterpreter, PDFResourceManager
from pdfminer.pdfpage import PDFPage
from utils.StorageClient import GoogleCloudStorageClient, AnnotatedDataStorageClient

# Mixins
class DocumentProcessingMixin:

    @staticmethod
    def extract_text_from_pdf(contents: io.BytesIO) -> list:
        resource_manager = PDFResourceManager()
        fake_file_handle = io.StringIO()
        converter = TextConverter(
            resource_manager, fake_file_handle, laparams=LAParams()
        )
        page_interpreter = PDFPageInterpreter(resource_manager, converter)
        for page in PDFPage.get_pages(contents, caching=True, check_extractable=True):
            page_interpreter.process_page(page)
        text = fake_file_handle.getvalue()

        converter.close()
        fake_file_handle.close()

        return text

    def get_text_splitter(self, document: Document) -> CharacterTextSplitter:
        options = PDFRetrivalOption(
            splitter=PDFSplitterOption(
                chunk_overlap=document.split_option.get("chunk_overlap", 0),
                chunk_size=document.split_option.get("chunk_size", 100),
            )
        )
        return CharacterTextSplitter.from_tiktoken_encoder(
            separator=" ",
            chunk_size=options.splitter.chunk_size,
            chunk_overlap=options.splitter.chunk_overlap,
        )

    def split_content(self, content: str) -> List[str]:
        return content.split("\f")

class DocumentHandler(ABC, DocumentProcessingMixin):

    @abstractmethod
    def fetch_content(self, document: Document) -> str:
        pass

    def generate_metadata(self, document: Document) -> dict:
        # By default, consider the URL as the source.
        return {
            "source": document.url
        }

    def process(self, document: Document, dataset: Dataset) -> List[Document]:
        content = self.fetch_content(document)
        document.content_size = sys.getsizeof(content)
        # Get the default metadata
        metadata = self.generate_metadata(document)
        docs = [Document(page_content=content_part, metadata=metadata.copy()) for content_part in self.split_content(content)]
        # Set the page size
        document.page_size = len(docs)
        for page_number, doc in enumerate(docs):
            doc.metadata["page_number"] = page_number
            doc.metadata["urn"] = f"{dataset.id}-{document.url}-{doc.metadata['page_number']}"

        logger.info(f"got documents: {len(docs)} while loading dataset {dataset.id}")
        return docs

class PDFHandler(DocumentHandler):
    
    def fetch_content(self, document: Document) -> str:
        storage_client=GoogleCloudStorageClient()
        pdf_content = storage_client.load(document.url)
        return self.extract_text_from_pdf(pdf_content)

class AnnotatedDataHandler(DocumentHandler):

    def fetch_content(self, document: Document) -> str:
        webhook_handler = AnnotatedDataStorageClient()
        return webhook_handler.load(document.uid)

    def split_content(self, content: str) -> List[str]:
        return [content]

    def generate_metadata(self, document: Document) -> dict:
        # For annotated data, the UID is the source.
        return {
            "source": document.uid
        }
class WordHandler(DocumentHandler):
    
    def fetch_content(self, document: Document) -> str:
        storage_client = GoogleCloudStorageClient()
        word_content = storage_client.load(document.url)
        word_doc = WordDocument(word_content)
        full_text = []
        for para in word_doc.paragraphs:
            full_text.append(para.text)
        
        return '\n'.join(full_text)
    
def load_and_split_documents(datasets: list[Dataset]):
    handlers = {
        "pdf": PDFHandler(),
        "annotated_data": AnnotatedDataHandler(),
    }

    docs = []
    for dataset in datasets:
        for document in dataset.documents:
            handler = handlers.get(document.type)
            if handler:
                processed_docs = handler.process(document, dataset) 
                docs.extend(processed_docs)
            else:
                # Handle unsupported document types
                logger.error(f"Document type {document.type} not supported")
                raise Exception("Document type not supported")

    return docs
